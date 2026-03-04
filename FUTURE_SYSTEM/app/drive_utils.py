from pathlib import Path
from typing import Dict, List, Optional, Tuple


def mount_drive_if_available(mount_point: str = "/content/drive") -> bool:
    """Mount Drive only when running inside Colab."""
    try:
        from google.colab import drive  # type: ignore
    except Exception:
        return False

    drive.mount(mount_point, force_remount=False)
    return True


def list_dir(path: str) -> List[str]:
    p = Path(path)
    if not p.exists():
        return []
    items = []
    for child in sorted(p.iterdir(), key=lambda x: x.name.lower()):
        suffix = "/" if child.is_dir() else ""
        items.append(f"{child.name}{suffix}")
    return items


def verify_shortcut_folder(
    folder_name: str = "FUTURE_SYSTEM", base_mydrive: str = "/content/drive/MyDrive"
) -> Tuple[bool, str, List[str]]:
    p = Path(base_mydrive) / folder_name
    exists = p.exists()
    return exists, str(p), list_dir(str(p)) if exists else []


def find_adapter_dir(search_roots: List[str]) -> Tuple[Optional[str], List[str]]:
    required_any = (
        "adapter_config.json",
        "adapter_model.safetensors",
        "tokenizer_config.json",
        "special_tokens_map.json",
    )
    for root in search_roots:
        p = Path(root)
        if not p.exists() or not p.is_dir():
            continue
        files = {f.name for f in p.iterdir() if f.is_file()}
        if any(name in files for name in required_any):
            return str(p), sorted(files)
    return None, []


def validate_project_and_adapter(
    project_root: str,
    adapter_candidates: List[str],
) -> Dict[str, object]:
    required_files = ["adapter_config.json", "adapter_model.safetensors"]
    optional_files = ["README.md"]

    result: Dict[str, object] = {
        "project_root": project_root,
        "project_exists": False,
        "project_items": [],
        "adapter_dir": "",
        "adapter_exists": False,
        "adapter_files": [],
        "required_status": {},
        "optional_status": {},
        "missing_required": [],
        "error": "",
    }

    root = Path(project_root)
    if not root.exists() or not root.is_dir():
        result["error"] = (
            "FUTURE_SYSTEM shortcut folder not found in MyDrive. "
            "Expected: /content/drive/MyDrive/FUTURE_SYSTEM"
        )
        return result

    result["project_exists"] = True
    result["project_items"] = list_dir(str(root))

    adapter_dir, files = find_adapter_dir(adapter_candidates)
    if not adapter_dir:
        result["error"] = (
            "Adapter folder not found under FUTURE_SYSTEM. "
            "Expected one of: my_finetuned_llm or my_finetined_llm."
        )
        return result

    result["adapter_exists"] = True
    result["adapter_dir"] = adapter_dir
    result["adapter_files"] = files

    file_set = set(files)
    req_status = {name: (name in file_set) for name in required_files}
    opt_status = {name: (name in file_set) for name in optional_files}
    missing_required = [name for name, ok in req_status.items() if not ok]

    result["required_status"] = req_status
    result["optional_status"] = opt_status
    result["missing_required"] = missing_required

    if missing_required:
        result["error"] = "Missing required adapter files: " + ", ".join(missing_required)

    return result


def _resolve_upload_path(upload_obj) -> Optional[str]:
    if upload_obj is None:
        return None
    if isinstance(upload_obj, str):
        return upload_obj
    if isinstance(upload_obj, dict):
        return upload_obj.get("name") or upload_obj.get("path")
    if hasattr(upload_obj, "name"):
        return getattr(upload_obj, "name")
    return None


def read_uploaded_text(upload_obj) -> str:
    path = _resolve_upload_path(upload_obj)
    if not path:
        return ""
    p = Path(path)
    if not p.exists():
        return ""

    ext = p.suffix.lower()
    if ext in {".txt", ".md", ".csv", ".json", ".py"}:
        return p.read_text(encoding="utf-8", errors="ignore")

    if ext == ".docx":
        try:
            from docx import Document  # type: ignore
        except Exception:
            return ""
        try:
            doc = Document(str(p))
            chunks = [para.text for para in doc.paragraphs if para.text]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            chunks.append(cell.text)
            return "\n".join(chunks).strip()
        except Exception:
            return ""

    if ext == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception:
            return ""
        try:
            reader = PdfReader(str(p))
            chunks = []
            for page in reader.pages:
                txt = page.extract_text() or ""
                if txt:
                    chunks.append(txt)
            return "\n".join(chunks).strip()
        except Exception:
            return ""

    # Fallback as text.
    return p.read_text(encoding="utf-8", errors="ignore")
